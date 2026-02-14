"""
VisioNova Document Parser
Extracts text from PDFs, DOCX, and TXT files with intelligent chunking.
"""
import os
import re
from typing import List, Dict, Optional, Tuple
from io import BytesIO
try:
    from PIL import Image
    import pytesseract
except ImportError:
    Image = None
    pytesseract = None


class DocumentParser:
    """
    Extract text from various document formats and chunk for analysis.
    
    Supported formats:
    - PDF (.pdf) using PyMuPDF + AI enhancement
    - Word (.docx) using python-docx + AI enhancement
    - Plain text (.txt)
    
    AI Enhancement:
    - Uses Llama 4 Scout via Groq for text cleanup and structuring
    - Fallback to raw extraction if AI fails
    """
    
    # Default chunk size (characters)
    DEFAULT_CHUNK_SIZE = 2000
    MIN_CHUNK_SIZE = 500
    MAX_CHUNK_SIZE = 5000
    
    # Supported file extensions
    SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.txt', '.doc'}
    
    def __init__(self, chunk_size: int = DEFAULT_CHUNK_SIZE, use_ai: bool = True):
        """Initialize parser with configurable chunk size and AI option."""
        self.chunk_size = max(self.MIN_CHUNK_SIZE, min(chunk_size, self.MAX_CHUNK_SIZE))
        self.use_ai = use_ai
        self.ai_extractor = None
        
        # Initialize AI client if requested
        if self.use_ai:
            try:
                from AI import AIDocumentExtractor
                self.ai_extractor = AIDocumentExtractor()
                print("[DocumentParser] AI extraction enabled")
            except Exception as e:
                print(f"[DocumentParser] AI initialization failed, using fallback: {e}")
                self.ai_extractor = None
    
    def parse_file(self, file_path: str) -> Dict:
        """
        Parse a file and extract text.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            dict with 'text', 'chunks', 'metadata', 'error'
        """
        if not os.path.exists(file_path):
            return {"error": f"File not found: {file_path}"}
        
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext not in self.SUPPORTED_EXTENSIONS:
            return {"error": f"Unsupported file format: {ext}"}
        
        try:
            if ext == '.pdf':
                text, metadata = self._extract_pdf(file_path)
            elif ext in ['.docx', '.doc']:
                text, metadata = self._extract_docx(file_path)
            else:  # .txt
                text, metadata = self._extract_txt(file_path)
            
            if not text or not text.strip():
                return {"error": "No text content found in file"}
            
            # Apply AI enhancement if available
            enhanced_text = self._enhance_with_ai(text, metadata) if self.ai_extractor else text
            
            chunks = self.chunk_text(enhanced_text)
            
            return {
                "text": enhanced_text,
                "chunks": chunks,
                "metadata": metadata,
                "ai_enhanced": self.ai_extractor is not None,
                "error": None
            }
            
        except Exception as e:
            return {"error": f"Failed to parse file: {str(e)}"}
    
    def parse_bytes(self, file_bytes: bytes, filename: str) -> Dict:
        """
        Parse file from bytes (for file uploads).
        
        Args:
            file_bytes: File content as bytes
            filename: Original filename (for extension detection)
            
        Returns:
            dict with 'text', 'chunks', 'metadata', 'error'
        """
        ext = os.path.splitext(filename)[1].lower()
        
        if ext not in self.SUPPORTED_EXTENSIONS:
            return {"error": f"Unsupported file format: {ext}"}
        
        try:
            if ext == '.pdf':
                text, metadata = self._extract_pdf_bytes(file_bytes)
            elif ext in ['.docx', '.doc']:
                text, metadata = self._extract_docx_bytes(file_bytes)
            else:  # .txt
                text = file_bytes.decode('utf-8', errors='ignore')
                metadata = {"format": "txt", "char_count": len(text)}
            
            if not text or not text.strip():
                return {"error": "No text content found in file"}
            
            # Apply AI enhancement if available
            enhanced_text = self._enhance_with_ai(text, metadata) if self.ai_extractor else text
            
            chunks = self.chunk_text(enhanced_text)
            
            return {
                "text": enhanced_text,
                "chunks": chunks,
                "metadata": metadata,
                "ai_enhanced": self.ai_extractor is not None,
                "error": None
            }
            
        except Exception as e:
            return {"error": f"Failed to parse file: {str(e)}"}
    
    def _extract_pdf(self, file_path: str) -> Tuple[str, Dict]:
        """Extract text from PDF file."""
        try:
            import fitz  # PyMuPDF
            # Helper to access Matrix for zooming
            global import_fitz_matrix
            import_fitz_matrix = fitz.Matrix
        except ImportError:
            raise ImportError("PyMuPDF not installed. Run: pip install PyMuPDF")
        
        doc = fitz.open(file_path)
        text_parts = []
        
        for page in doc:
            # Use sort=True for natural reading order (top-left to bottom-right)
            text_parts.append(page.get_text(sort=True))
        
        # doc.close() defered to end for OCR capability ->
        
        full_text = "\n\n".join(text_parts)
        
        # Check for scanned PDF (low text density)
        # Avg chars per page < 50 usually means it's an image scan or very sparse
        avg_chars = len(full_text) / len(text_parts) if text_parts else 0
        
        if avg_chars < 50 and pytesseract:
            print("[DocumentParser] Low text density detected. Attempting OCR...")
            ocr_text = self._extract_pdf_ocr(doc)
            if len(ocr_text) > len(full_text):
                full_text = ocr_text
                print(f"[DocumentParser] OCR successful (extracted {len(full_text)} chars)")
        
        doc.close()
        
        metadata = {
            "format": "pdf",
            "pages": len(text_parts),
            "char_count": len(full_text),
            "ocr_applied": avg_chars < 50
        }
        
        return full_text, metadata
    
    def _extract_pdf_ocr(self, doc) -> str:
        """Extract text from PDF using OCR (images)."""
        text_parts = []
        
        for i, page in enumerate(doc):
            try:
                # Render page to image
                # Zoom = 2.0 for better quality (300 DPI approx)
                pix = page.get_pixmap(matrix=import_fitz_matrix(2.0, 2.0))
                img_data = pix.tobytes("png")
                
                # Create PIL Image
                img = Image.open(BytesIO(img_data))
                
                # Run OCR
                text = pytesseract.image_to_string(img)
                text_parts.append(text)
                
            except Exception as e:
                print(f"OCR failed for page {i+1}: {e}")
                
        return "\n\n".join(text_parts)
    
    def _extract_pdf_bytes(self, file_bytes: bytes) -> Tuple[str, Dict]:
        """Extract text from PDF bytes."""
        try:
            import fitz  # PyMuPDF
        except ImportError:
            raise ImportError("PyMuPDF not installed. Run: pip install PyMuPDF")
        
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        text_parts = []
        
        for page in doc:
            # Use sort=True for natural reading order
            text_parts.append(page.get_text(sort=True))
        
        # Join text parts first so full_text is defined before use
        full_text = "\n\n".join(text_parts)
        
        # Check for scanned PDF (low text density)
        avg_chars = len(full_text) / len(text_parts) if text_parts else 0
        
        if avg_chars < 50 and pytesseract:
            print("[DocumentParser] Low text density detected (bytes). Attempting OCR...")
            ocr_text = self._extract_pdf_ocr(doc)
            if len(ocr_text) > len(full_text):
                full_text = ocr_text
                print(f"[DocumentParser] OCR successful (extracted {len(full_text)} chars)")
        
        doc.close()
        
        metadata = {
            "format": "pdf",
            "pages": len(text_parts),
            "char_count": len(full_text),
            "ocr_applied": avg_chars < 50
        }
        
        return full_text, metadata
    
    def _extract_docx(self, file_path: str) -> Tuple[str, Dict]:
        """Extract text from DOCX file."""
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx not installed. Run: pip install python-docx")
        
        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        
        full_text = "\n\n".join(paragraphs)
        metadata = {
            "format": "docx",
            "paragraphs": len(paragraphs),
            "char_count": len(full_text)
        }
        
        return full_text, metadata
    
    def _extract_docx_bytes(self, file_bytes: bytes) -> Tuple[str, Dict]:
        """Extract text from DOCX bytes."""
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx not installed. Run: pip install python-docx")
        
        doc = Document(BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        
        full_text = "\n\n".join(paragraphs)
        metadata = {
            "format": "docx",
            "paragraphs": len(paragraphs),
            "char_count": len(full_text)
        }
        
        return full_text, metadata
    
    def _extract_txt(self, file_path: str) -> Tuple[str, Dict]:
        """Extract text from TXT file."""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            text = f.read()
        
        metadata = {
            "format": "txt",
            "char_count": len(text)
        }
        
        return text, metadata
    
    def chunk_text(self, text: str) -> List[Dict]:
        """
        Split text into chunks at sentence boundaries.
        
        Args:
            text: Full text to chunk
            
        Returns:
            List of chunk dicts with text, start, end positions
        """
        if len(text) <= self.chunk_size:
            return [{
                "text": text,
                "start": 0,
                "end": len(text),
                "index": 0
            }]
        
        # Split into sentences
        sentences = re.split(r'(?<=[.!?])\s+', text)
        
        chunks = []
        current_chunk = ""
        current_start = 0
        chunk_index = 0
        
        for sentence in sentences:
            # If adding this sentence exceeds chunk size, save current chunk
            if len(current_chunk) + len(sentence) > self.chunk_size and current_chunk:
                chunks.append({
                    "text": current_chunk.strip(),
                    "start": current_start,
                    "end": current_start + len(current_chunk),
                    "index": chunk_index
                })
                chunk_index += 1
                current_start += len(current_chunk)
                current_chunk = ""
            
            current_chunk += sentence + " "
        
        # Don't forget the last chunk
        if current_chunk.strip():
            chunks.append({
                "text": current_chunk.strip(),
                "start": current_start,
                "end": current_start + len(current_chunk),
                "index": chunk_index
            })
        
        return chunks
    
    @staticmethod
    def is_supported(filename: str) -> bool:
        """Check if file format is supported."""
        ext = os.path.splitext(filename)[1].lower()
        return ext in DocumentParser.SUPPORTED_EXTENSIONS
    
    @staticmethod
    def get_supported_formats() -> List[str]:
        """Get list of supported file extensions."""
        return list(DocumentParser.SUPPORTED_EXTENSIONS)
    
    def _enhance_with_ai(self, text: str, metadata: Dict) -> str:
        """
        Enhance extracted text using AI for better structure and cleanup.
        
        Args:
            text: Raw extracted text
            metadata: Document metadata (format, pages, etc.)
            
        Returns:
            AI-enhanced text, or original if AI fails
        """
        if not self.ai_extractor:
            return text
        
        try:
            # Only apply AI enhancement for documents that might benefit
            # (longer documents with potential formatting issues)
            if len(text) < 500:  # Skip AI for very short texts
                return text
            
            print("[DocumentParser] Applying AI enhancement...")
            format_type = metadata.get('format', 'unknown')
            pages = metadata.get('pages', metadata.get('paragraphs', 0))
            
            context = f"Document format: {format_type.upper()}, Pages/Sections: {pages}"
            enhanced = self.ai_extractor.extract_from_pages([text])
            
            print(f"[DocumentParser] AI enhancement complete (original: {len(text)} chars, enhanced: {len(enhanced)} chars)")
            return enhanced if enhanced and len(enhanced) > 100 else text
            
        except Exception as e:
            print(f"[DocumentParser] AI enhancement failed, using raw text: {e}")
            return text


if __name__ == "__main__":
    # Test the parser
    parser = DocumentParser()
    
    # Test chunking
    sample_text = """
    This is the first sentence. Here is the second sentence with more words.
    And now a third sentence that follows. The fourth sentence continues the thought.
    We have a fifth sentence here. Sixth sentence adds more content.
    Seventh sentence keeps going. Eighth provides additional text.
    Ninth sentence wraps up this paragraph. Tenth sentence concludes.
    """ * 10
    
    chunks = parser.chunk_text(sample_text)
    print(f"Text length: {len(sample_text)}")
    print(f"Number of chunks: {len(chunks)}")
    for i, chunk in enumerate(chunks[:3]):
        print(f"Chunk {i}: {len(chunk['text'])} chars")
